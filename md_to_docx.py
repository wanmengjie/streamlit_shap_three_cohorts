# -*- coding: utf-8 -*-
"""Convert PAPER_DRAFT_完整版.md to Word (.docx)."""
import os
from docx import Document
from docx.shared import Pt

def main():
    base = os.path.dirname(os.path.abspath(__file__))
    md_path = os.path.join(base, "PAPER_DRAFT_完整版.md")
    out_path = os.path.join(base, "PAPER_DRAFT_Full_Paper.docx")

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    i = 0
    table_rows = []

    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=0)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=2)
        elif stripped == "---":
            pass
        elif stripped.startswith("|") and "|" in stripped[1:]:
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if cells:
                table_rows.append(cells)
            i += 1
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].split("|")[1:-1]]
                if cells:
                    table_rows.append(cells)
                i += 1
            if table_rows:
                ncols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=ncols)
                table.style = "Table Grid"
                for ri, row in enumerate(table_rows):
                    for ci, cell in enumerate(row):
                        if ci < ncols:
                            table.rows[ri].cells[ci].text = cell
                doc.add_paragraph()
                table_rows = []
            continue
        elif stripped:
            p = doc.add_paragraph()
            if stripped.startswith("**") and ":**" in stripped:
                idx = stripped.find(":**")
                r1 = p.add_run(stripped[:idx+2])
                r1.bold = True
                p.add_run(stripped[idx+2:])
            elif stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
                r = p.add_run(stripped[1:-1])
                r.italic = True
            else:
                p.add_run(stripped)
        i += 1

    doc.save(out_path)
    print("Saved:", out_path)
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop", "PAPER_DRAFT_Full_Paper.docx")
    if desktop_path != out_path:
        doc.save(desktop_path)
        print("Also saved to Desktop:", desktop_path)

if __name__ == "__main__":
    main()
